"""
Parser de currículo com IA.

Extrai texto de PDF/DOCX e usa OpenAI para estruturar o perfil do candidato.
"""

from __future__ import annotations

import json
import os
import shutil
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument
from loguru import logger

from ai.client import get_ai_client
from ai.schemas import RESUME_SCHEMA
from config.prompts import RESUME_PARSER_SYSTEM, RESUME_PARSER_USER
from config.settings import CV_UPLOADS_DIR
from core.database import ResumeRecord, get_session
from core.exceptions import (
    ResumeFileTooLargeError,
    ResumeParsingError,
    ResumeUnsupportedFormatError,
)
from core.models import ResumeProfile
from utils.helpers import compute_file_hash, get_file_size_mb, sanitize_text


def extract_text_from_pdf(file_path: str) -> str:
    """Extrai texto de um arquivo PDF."""
    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        if not text_parts:
            raise ResumeParsingError("PDF sem texto extraível.")
        return "\n\n".join(text_parts)
    except ResumeParsingError:
        raise
    except Exception as e:
        raise ResumeParsingError(f"Erro ao ler PDF: {e}") from e


def extract_text_from_docx(file_path: str) -> str:
    """Extrai texto de um arquivo DOCX."""
    try:
        doc = DocxDocument(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        if not parts:
            raise ResumeParsingError("DOCX vazio.")
        return "\n".join(parts)
    except ResumeParsingError:
        raise
    except Exception as e:
        raise ResumeParsingError(f"Erro ao ler DOCX: {e}") from e


def extract_text(file_path: str) -> str:
    """Extrai texto de PDF ou DOCX e sanitiza."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        text = extract_text_from_docx(file_path)
    else:
        raise ResumeUnsupportedFormatError(f"Formato não suportado: {ext}")
    return sanitize_text(text)


def validate_resume_file(file_path: str) -> None:
    """Valida arquivo de currículo (existência, tamanho, formato)."""
    if not os.path.exists(file_path):
        raise ResumeParsingError(f"Arquivo não encontrado: {file_path}")
    if get_file_size_mb(file_path) > 10:
        raise ResumeFileTooLargeError("Arquivo > 10MB.")
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in (".pdf", ".docx", ".doc"):
        raise ResumeUnsupportedFormatError(f"Formato não suportado: {ext}")


def get_cached_profile(file_hash: str) -> Optional[ResumeProfile]:
    """Busca perfil no cache pelo hash do arquivo."""
    try:
        session = get_session()
        record = session.query(ResumeRecord).filter_by(file_hash=file_hash).first()
        session.close()
        if record and record.parsed_json:
            return ResumeProfile(**json.loads(record.parsed_json))
    except Exception as e:
        logger.warning(f"Erro ao buscar cache de CV: {e}")
    return None


def save_profile_cache(file_path: str, file_hash: str, profile_data: dict) -> None:
    """Salva o perfil analisado no cache."""
    try:
        session = get_session()
        existing = session.query(ResumeRecord).filter_by(file_hash=file_hash).first()
        if existing:
            existing.parsed_json = json.dumps(profile_data, ensure_ascii=False)
        else:
            session.add(ResumeRecord(
                file_path=file_path,
                file_name=os.path.basename(file_path),
                file_hash=file_hash,
                parsed_json=json.dumps(profile_data, ensure_ascii=False),
            ))
        session.commit()
        session.close()
    except Exception as e:
        logger.warning(f"Erro ao salvar cache de CV: {e}")


def copy_resume_to_uploads(file_path: str) -> str:
    """Copia o currículo para o diretório de uploads."""
    os.makedirs(CV_UPLOADS_DIR, exist_ok=True)
    dest = os.path.join(CV_UPLOADS_DIR, os.path.basename(file_path))
    if os.path.abspath(file_path) != os.path.abspath(dest):
        shutil.copy2(file_path, dest)
    return dest


async def parse_resume(file_path: str) -> ResumeProfile:
    """Processa currículo: extrai texto, analisa com IA, retorna perfil. Usa cache."""
    validate_resume_file(file_path)
    file_hash = compute_file_hash(file_path)

    cached = get_cached_profile(file_hash)
    if cached:
        logger.info("CV encontrado no cache")
        return cached

    logger.info("Extraindo texto do currículo...")
    resume_text = extract_text(file_path)
    if len(resume_text.strip()) < 50:
        raise ResumeParsingError("Texto extraído muito curto.")

    saved_path = copy_resume_to_uploads(file_path)

    logger.info("Analisando currículo com IA...")
    client = get_ai_client()
    prompt = RESUME_PARSER_USER.format(resume_text=resume_text[:8000])
    result = await client.chat_completion(
        system_prompt=RESUME_PARSER_SYSTEM,
        user_prompt=prompt,
        response_schema=RESUME_SCHEMA,
        call_type="resume_parse",
    )

    save_profile_cache(saved_path, file_hash, result)
    profile = ResumeProfile(**result)
    logger.info(f"CV analisado: {profile.nome} — {len(profile.skills_tecnicas)} skills")
    return profile
